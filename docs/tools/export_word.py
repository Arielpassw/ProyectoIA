"""Exporta la documentación Markdown del proyecto a documentos Word editables."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
DOCS_DIR = ROOT / "docs"
OUTPUT_DIR = DOCS_DIR / "word"
ASSETS_DIR = DOCS_DIR / "assets"

DOCUMENTS = {
    "INFORME_FINAL.md": ("INFORME FINAL", "Informe final del proyecto"),
    "MANUAL_TECNICO.md": ("MANUAL TÉCNICO", "Manual técnico"),
    "DOCUMENTACION_SISTEMA.md": (
        "DOCUMENTACIÓN DEL SISTEMA",
        "Documentación del sistema",
    ),
    "API.md": ("REFERENCIA DE LA API", "Documentación de la API REST"),
    "GUIA_DESARROLLADORES.md": (
        "GUÍA PARA DESARROLLADORES",
        "Guía paso a paso para nuevos desarrolladores",
    ),
}

NAVY = "17365D"
BLUE = "1F4E78"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "E7E6E6"
DARK_GRAY = "404040"


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def add_bottom_border(paragraph, color=NAVY, size="12") -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    border = OxmlElement("w:bottom")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), "3")
    border.set(qn("w:color"), color)
    borders.append(border)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Página ")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in (
        ("Title", 22, NAVY),
        ("Heading 1", 16, NAVY),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11, DARK_GRAY),
    ):
        style = document.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_gray_bar(document: Document) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(16)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    cell.height = Cm(0.28)
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)


def add_cover(document: Document, document_kind: str, title: str) -> None:
    header = document.add_table(rows=1, cols=3)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    header.autofit = False
    widths = (Cm(2.4), Cm(11.4), Cm(2.2))
    for cell, width in zip(header.rows[0].cells, widths):
        cell.width = width
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    left = header.cell(0, 0).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left.add_run().add_picture(str(ASSETS_DIR / "epn-logo.png"), width=Cm(1.55))

    center = header.cell(0, 1)
    center.text = ""
    p = center.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ESCUELA POLITÉCNICA NACIONAL")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(16)
    p2 = center.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("ESCUELA DE FORMACIÓN DE TECNÓLOGOS")
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    add_bottom_border(p2)
    p3 = center.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p3.add_run("TECNOLOGÍA SUPERIOR EN DESARROLLO DE SOFTWARE")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)

    right = header.cell(0, 2).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    right.add_run().add_picture(str(ASSETS_DIR / "buho-epn.png"), width=Cm(1.15))

    document.add_paragraph()
    add_gray_bar(document)

    info = document.add_table(rows=4, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.autofit = False
    info.columns[0].width = Cm(6.2)
    info.columns[1].width = Cm(9.8)
    values = (
        ("ASIGNATURA:", "Fundamentos de Inteligencia Artificial"),
        ("PROFESOR:", "Ing. Vanessa Guevara"),
        ("PERÍODO ACADÉMICO:", "2026-A"),
        ("ESTUDIANTE:", "Emilia Tana"),
    )
    for row, (label, value) in zip(info.rows, values):
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].bold = True
        for cell in row.cells:
            set_cell_margins(cell, top=30, bottom=30)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)

    add_gray_bar(document)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    r = p.add_run("PROYECTO 2DO BIMESTRE")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(16)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run(document_kind)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("TÍTULO:")
    r.font.name = "Times New Roman"
    r.font.size = Pt(14)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Clasificador de prendas con inteligencia artificial")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(15)

    screenshot = DOCS_DIR / "capturas" / "interfaz-principal.png"
    if screenshot.exists():
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        p.add_run().add_picture(str(screenshot), width=Cm(10.2))

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("Emilia Tana")
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Quito – Ecuador\n2026")
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)

    document.add_page_break()

    p = document.add_paragraph(title, style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)


INLINE_PATTERN = re.compile(
    r"(\*\*.+?\*\*|`.+?`|\[[^]]+\]\([^)]+\)|\*[^*]+\*)"
)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend((color, underline))
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.extend((properties, text_element))
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, text: str) -> None:
    position = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
        elif token.startswith("["):
            link = re.match(r"\[([^]]+)\]\(([^)]+)\)", token)
            if link:
                add_hyperlink(paragraph, link.group(1), link.group(2))
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def add_table(document: Document, rows: list[str]) -> None:
    parsed = [[cell.strip() for cell in row.strip().strip("|").split("|")] for row in rows]
    if len(parsed) > 1 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in parsed[1]):
        parsed.pop(1)
    if not parsed:
        return
    width = max(len(row) for row in parsed)
    table = document.add_table(rows=len(parsed), cols=width)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_index, values in enumerate(parsed):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            cell.text = ""
            add_inline(cell.paragraphs[0], value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, NAVY)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
        if row_index == 0:
            set_repeat_table_header(table.rows[0])
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_markdown(document: Document, source: Path) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []
    paragraph_lines: list[str] = []
    skip_academic_front = source.name == "INFORME_FINAL.md"
    content_started = not skip_academic_front

    def flush_paragraph() -> None:
        if not paragraph_lines:
            return
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_inline(paragraph, " ".join(part.strip() for part in paragraph_lines))
        paragraph_lines.clear()

    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()

        if skip_academic_front and not content_started:
            if stripped.startswith("## 1."):
                content_started = True
            else:
                index += 1
                continue

        if stripped.startswith("```"):
            flush_paragraph()
            if in_code:
                paragraph = document.add_paragraph()
                paragraph.style = document.styles["Normal"]
                paragraph.paragraph_format.left_indent = Cm(0.5)
                paragraph.paragraph_format.right_indent = Cm(0.5)
                paragraph.paragraph_format.space_before = Pt(4)
                paragraph.paragraph_format.space_after = Pt(8)
                run = paragraph.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(8.5)
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "F2F2F2")
                paragraph._p.get_or_add_pPr().append(shading)
                code_lines.clear()
            in_code = not in_code
            index += 1
            continue
        if in_code:
            code_lines.append(raw)
            index += 1
            continue

        if not stripped:
            flush_paragraph()
            index += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            flush_paragraph()
            table_rows = []
            while index < len(lines):
                candidate = lines[index].strip()
                if not (candidate.startswith("|") and candidate.endswith("|")):
                    break
                table_rows.append(candidate)
                index += 1
            add_table(document, table_rows)
            continue

        image = re.fullmatch(r"!\[([^]]*)\]\(([^)]+)\)", stripped)
        if image:
            flush_paragraph()
            image_path = (source.parent / image.group(2)).resolve()
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if image_path.exists():
                paragraph.add_run().add_picture(str(image_path), width=Cm(15.5))
                caption = document.add_paragraph(image.group(1))
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if caption.runs:
                    caption.runs[0].italic = True
                    caption.runs[0].font.size = Pt(9)
            else:
                paragraph.add_run(f"[Imagen pendiente: {image.group(1)}]").italic = True
            index += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            level = len(heading.group(1))
            # El H1 ya está representado en la portada y el título interior.
            if level > 1:
                paragraph = document.add_paragraph(
                    heading.group(2), style=f"Heading {min(level - 1, 3)}"
                )
                paragraph.paragraph_format.keep_with_next = True
            index += 1
            continue

        if re.fullmatch(r"-{3,}", stripped):
            flush_paragraph()
            paragraph = document.add_paragraph()
            add_bottom_border(paragraph, color=LIGHT_GRAY, size="8")
            index += 1
            continue

        unordered = re.match(r"^[-*]\s+(.+)$", stripped)
        ordered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if unordered or ordered:
            flush_paragraph()
            value = (unordered or ordered).group(1)
            paragraph = document.add_paragraph(
                style="List Bullet" if unordered else "List Number"
            )
            add_inline(paragraph, value)
            index += 1
            continue

        if stripped.startswith(">"):
            flush_paragraph()
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.7)
            paragraph.paragraph_format.right_indent = Cm(0.4)
            run = paragraph.add_run(stripped.lstrip("> "))
            run.italic = True
            run.font.color.rgb = RGBColor.from_string(DARK_GRAY)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), LIGHT_BLUE)
            paragraph._p.get_or_add_pPr().append(shading)
            index += 1
            continue

        paragraph_lines.append(stripped.rstrip("  "))
        index += 1

    flush_paragraph()


def add_headers_and_footers(document: Document, title: str) -> None:
    for section in document.sections:
        header = section.header.paragraphs[0]
        header.text = "EPN · ESFOT · Fundamentos de Inteligencia Artificial"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header.runs:
            run.font.name = "Arial"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string("7F7F7F")
        add_page_number(section.footer.paragraphs[0])


def export(source_name: str, document_kind: str, title: str) -> Path:
    document = Document()
    configure_document(document)
    add_cover(document, document_kind, title)
    add_markdown(document, DOCS_DIR / source_name)
    add_headers_and_footers(document, title)
    document.core_properties.title = title
    document.core_properties.author = "Emilia Tana"
    document.core_properties.subject = "Proyecto 2do Bimestre"
    output = OUTPUT_DIR / f"{Path(source_name).stem}.docx"
    document.save(output)
    return output


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    for source, (kind, title) in DOCUMENTS.items():
        output = export(source, kind, title)
        print(f"Creado: {output.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
